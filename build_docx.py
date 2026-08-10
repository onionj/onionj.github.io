"""
build_docx.py — Generate an ATS-friendly Word (.docx) CV from index.md + _config.yml.

Usage:
    pip install python-docx pyyaml
    python build_docx.py
"""

import re
import yaml
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Config knobs ──────────────────────────────────────────────────────────────
MARGIN_CM = 1.5
MAX_BULLETS_PER_ROLE = 0  # 0 = no limit; set to 3 for a tight one-pager

# Same palette as header.tex / the site theme, so all three outputs match.
# Hierarchy comes from size and spacing; headings are never lighter than the
# body text under them. Grey is reserved for metadata (dates), not headings.
ACCENT_COLOR = RGBColor(0x10, 0x78, 0x78)  # teal — name, section titles, links
DARK_COLOR = RGBColor(0x15, 0x20, 0x2B)    # near-black — all headings and body
META_COLOR = RGBColor(0x5A, 0x66, 0x75)    # grey — dates and the stack line

# '### ' headings that are really top-level sections, not employers.
TOP_LEVEL_SUBHEADINGS = {
    'open source',
    'technical skills',
    'soft skills',
    'education',
    'languages',
}

# ── Load sources ──────────────────────────────────────────────────────────────
with open("_config.yml") as f:
    config = yaml.safe_load(f)

with open("index.md") as f:
    body = f.read()

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, text, url):
    """Add a real clickable hyperlink to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def strip_bold_markers(text):
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)


def add_inline_formatting(paragraph, text):
    """Parse inline **bold** and [text](url) and add runs to the paragraph."""
    pattern = re.compile(
        r'\[(.+?)\]\((.+?)\)'       # link (check first — may contain **)
        r'|\*\*(.+?)\*\*'           # bold
        r'|([^*\[]+)'               # plain text
        r'|(.)'                      # fallback single char
    )
    for m in pattern.finditer(text):
        if m.group(1):
            link_text = strip_bold_markers(m.group(1))
            add_hyperlink(paragraph, link_text, m.group(2))
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):
            paragraph.add_run(m.group(4))
        elif m.group(5):
            paragraph.add_run(m.group(5))


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    set_paragraph_spacing(p, before=14, after=3)
    # accent rule under the heading, mirroring the PDF
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '8',
        qn('w:space'): '2',
        qn('w:color'): '107878',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_role_heading(doc, company, dates):
    """Employer + dates — the anchor of each job block, so it must dominate."""
    p = doc.add_paragraph()
    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_COLOR
    if dates:
        run_sep = p.add_run("  |  ")
        run_sep.font.size = Pt(10)
        run_sep.font.color.rgb = META_COLOR
        run_d = p.add_run(dates)
        run_d.font.size = Pt(10)
        run_d.font.color.rgb = META_COLOR
    set_paragraph_spacing(p, before=12, after=2)
    return p


def add_stack_line(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = META_COLOR
    set_paragraph_spacing(p, before=0, after=1)
    return p


def add_sub_heading(doc, text):
    """Project name — subordinate through size and spacing, not a lighter colour."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_COLOR
    set_paragraph_spacing(p, before=6, after=1)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    add_inline_formatting(p, text.lstrip('* ').lstrip('- '))
    for run in p.runs:
        run.font.size = Pt(9.5)
    set_paragraph_spacing(p, before=0, after=1)
    return p


def add_plain(doc, text):
    p = doc.add_paragraph()
    add_inline_formatting(p, text)
    for run in p.runs:
        run.font.size = Pt(9.5)
    set_paragraph_spacing(p, before=0, after=1)
    return p


# ── Parse index.md ────────────────────────────────────────────────────────────

sections = {}
current_section = None
current_lines = []

for line in body.split('\n'):
    heading = None
    if line.startswith('## '):
        heading = line.lstrip('#').strip().rstrip(':')
    elif line.startswith('### '):
        # Promote '### Technical Skills' and friends: they are peers of
        # 'Work Experience', not employers nested inside it.
        candidate = line.lstrip('#').strip().rstrip(':')
        if candidate.lower() in TOP_LEVEL_SUBHEADINGS:
            heading = candidate

    if heading:
        if current_section:
            sections[current_section] = current_lines
        current_section = heading
        current_lines = []
    elif current_section is not None:
        current_lines.append(line)

if current_section:
    sections[current_section] = current_lines

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# margins
for section in doc.sections:
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

# default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# ── Header: name + contacts ──────────────────────────────────────────────────

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = name_p.add_run(config['full_name'])
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = ACCENT_COLOR
set_paragraph_spacing(name_p, before=0, after=3)

contacts = []
if config.get('gmail'):
    contacts.append(config['gmail_url'])
if config.get('linkedin'):
    contacts.append(config['linkedin_url'].removeprefix('https://www.'))
if config.get('github'):
    contacts.append(config['github_url'].removeprefix('https://'))
if config.get('phone') and config.get('phone_number'):
    contacts.append(config['phone_number'])

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = contact_p.add_run('  |  '.join(contacts))
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_COLOR
set_paragraph_spacing(contact_p, before=0, after=4)

# ── Sections ──────────────────────────────────────────────────────────────────

SECTION_ORDER = [
    'Professional Summary',
    'Work Experience',
    'Technical Skills',
    'Soft Skills',
    'Open Source',
    'Education',
    'Languages',
]

for section_name in SECTION_ORDER:
    lines = sections.get(section_name)
    if not lines:
        for key in sections:
            if section_name.lower() in key.lower():
                lines = sections[key]
                break
    if not lines:
        continue

    add_section_heading(doc, section_name)

    if section_name == 'Professional Summary':
        text = ' '.join(l.strip() for l in lines if l.strip())
        add_plain(doc, text)
        continue

    if section_name == 'Work Experience':
        bullet_count = 0
        current_company = None
        for line in lines:
            if line.startswith('### '):
                # Require spaces around the separating dash, so hyphenated
                # company names ("Mehr-e Pars ICT") are not split at the hyphen.
                match = re.match(r'###\s+(.+?)\s+-\s+(.+)', line)
                if match:
                    current_company = match.group(1).strip()
                    dates = match.group(2).strip()
                    add_role_heading(doc, current_company, dates)
                    bullet_count = 0
                else:
                    heading_text = line.lstrip('#').strip()
                    add_role_heading(doc, heading_text, '')
                    bullet_count = 0
            elif line.startswith('#### '):
                add_sub_heading(doc, line.lstrip('#').strip())
            elif line.startswith('- Stack:'):
                add_stack_line(doc, line.lstrip('- ').strip())
            elif line.startswith('* ') or line.startswith('- '):
                if MAX_BULLETS_PER_ROLE and bullet_count >= MAX_BULLETS_PER_ROLE:
                    continue
                add_bullet(doc, line)
                bullet_count += 1
        continue

    if section_name in ('Technical Skills', 'Soft Skills', 'Languages', 'Education'):
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('* ') or stripped.startswith('- '):
                add_bullet(doc, stripped)
        continue

    if 'open source' in section_name.lower():
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('- '):
                add_bullet(doc, stripped)
        continue

    for line in lines:
        stripped = line.strip()
        if stripped:
            add_plain(doc, stripped)

# ── Save ──────────────────────────────────────────────────────────────────────

output_name = f"{config['full_name'].replace(' ', '_').lower()}_cv.docx"
doc.save(output_name)
print(f"Created: {output_name}")
